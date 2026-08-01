import os
import io
import re
from flask import Flask, request, jsonify
from flask_cors import CORS
from pypdf import PdfReader

try:
    import docx
except ImportError:
    docx = None

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FRONTEND_DIR = os.path.abspath(os.path.join(BASE_DIR, "../frontend"))

app = Flask(__name__, static_folder=FRONTEND_DIR, static_url_path="")
CORS(app)

PORT = 3000

# ---------------------------------------------------------------------------
# SKILLS DICTIONARY & CATEGORIES
# ---------------------------------------------------------------------------
SKILL_CATEGORIES = {
    "programming_languages": [
        "python", "javascript", "typescript", "java", "c++", "c#", "ruby", "golang",
        "go", "php", "rust", "swift", "kotlin", "sql", "html", "css", "bash", "shell", "r", "matlab"
    ],
    "frameworks": [
        "react", "vue", "angular", "next.js", "next", "nuxt", "svelte", "django",
        "flask", "fastapi", "spring boot", "spring", "express", "nest.js", "nestjs",
        "laravel", "rails", "jquery", "tailwind", "bootstrap", "pytorch", "tensorflow", "keras"
    ],
    "databases": [
        "postgresql", "postgres", "mysql", "mongodb", "redis", "sqlite", "firebase",
        "oracle", "elasticsearch", "dynamodb", "mariadb", "cassandra"
    ],
    "cloud_and_devops": [
        "aws", "gcp", "azure", "docker", "kubernetes", "k8s", "terraform", "ansible",
        "jenkins", "ci/cd", "circleci", "github actions", "linux", "nginx", "apache"
    ],
    "tools_and_libraries": [
        "git", "github", "gitlab", "npm", "webpack", "vite", "postman", "pandas",
        "numpy", "scikit-learn", "sklearn", "matplotlib", "seaborn", "graphql", "rest api"
    ],
    "ai_tools": [
        "openai", "gemini", "langchain", "llama", "transformers", "huggingface",
        "opencv", "rag", "vector database", "chromadb", "pinecone"
    ],
    "soft_skills": [
        "communication", "leadership", "teamwork", "problem solving", "critical thinking",
        "agile", "scrum", "collaboration", "time management", "adaptability", "mentorship", "presentation"
    ]
}

# Role Benchmark Expectations
ROLE_EXPECTED_SKILLS = {
    "Frontend Developer": [
        "html", "css", "javascript", "typescript", "react", "vue", "angular", "next",
        "tailwind", "bootstrap", "webpack", "vite", "git", "rest api"
    ],
    "Backend Developer": [
        "python", "java", "golang", "go", "php", "node", "express", "django", "flask",
        "fastapi", "spring", "sql", "postgresql", "mysql", "mongodb", "redis", "docker", "rest api"
    ],
    "Java Developer": [
        "java", "spring boot", "spring", "hibernate", "maven", "gradle", "microservices",
        "sql", "postgresql", "mysql", "junit", "docker", "git"
    ],
    "Python Developer": [
        "python", "django", "flask", "fastapi", "sql", "postgresql", "mongodb", "redis",
        "celery", "pytest", "docker", "git", "rest api"
    ],
    "Full Stack Developer": [
        "javascript", "typescript", "react", "node", "express", "sql", "postgresql",
        "mongodb", "html", "css", "git", "docker", "aws", "rest api"
    ],
    "Data Analyst": [
        "python", "sql", "excel", "tableau", "powerbi", "pandas", "numpy", "statistics",
        "analytics", "reporting", "dashboards", "postgresql"
    ],
    "ML Engineer": [
        "python", "pytorch", "tensorflow", "keras", "sklearn", "pandas", "numpy",
        "machine learning", "deep learning", "nlp", "models", "ai"
    ],
    "AI Engineer": [
        "python", "openai", "langchain", "transformers", "huggingface", "pytorch",
        "rag", "vector database", "machine learning", "deep learning", "fastapi"
    ],
    "Data Scientist": [
        "python", "r", "sql", "pandas", "numpy", "scikit-learn", "statistics", "machine learning",
        "data visualization", "data mining"
    ],
    "DevOps Engineer": [
        "docker", "kubernetes", "k8s", "aws", "terraform", "ansible", "jenkins", "ci/cd",
        "linux", "bash", "python", "git"
    ],
    "Cloud Engineer": [
        "aws", "gcp", "azure", "terraform", "cloud", "docker", "kubernetes", "linux",
        "networking", "security", "python"
    ],
    "Android Developer": [
        "kotlin", "java", "android sdk", "jetpack compose", "retrofit", "room", "sqlite",
        "git", "rest api"
    ],
    "Software Engineer": [
        "python", "java", "c++", "javascript", "sql", "git", "data structures",
        "algorithms", "object-oriented", "testing"
    ],
    "Cyber Security": [
        "security", "networking", "firewalls", "penetration testing", "vulnerability assessment",
        "linux", "python", "siem", "compliance", "cryptography"
    ]
}

ALL_KNOWN_SKILLS = set()
for category, skills in SKILL_CATEGORIES.items():
    ALL_KNOWN_SKILLS.update(skills)

@app.after_request
def add_cors_headers(response):
    response.headers["Access-Control-Allow-Origin"] = "*"
    response.headers["Access-Control-Allow-Headers"] = "Content-Type, Authorization"
    response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
    return response

@app.route("/")
def index():
    return app.send_static_file("index.html")

@app.route("/analyze", methods=["POST", "OPTIONS"])
def analyze_resume():
    if request.method == "OPTIONS":
        return "", 204

    try:
        if "resume" not in request.files:
            return jsonify({"error": "No resume file provided in request."}), 400

        file = request.files["resume"]
        if not file or file.filename.strip() == "":
            return jsonify({"error": "No file selected for upload."}), 400

        filename = file.filename.lower()
        if not (filename.endswith(".txt") or filename.endswith(".pdf") or filename.endswith(".docx")):
            return jsonify({"error": "Unsupported file format. Please upload a .pdf, .docx, or .txt file."}), 400

        raw_bytes = file.read()
        if not raw_bytes or len(raw_bytes.strip()) == 0:
            return jsonify({"error": "Uploaded file is empty (0 bytes)."}), 400

        resume_text = ""

        # Extract Text
        if filename.endswith(".txt"):
            try:
                resume_text = raw_bytes.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    resume_text = raw_bytes.decode("latin-1")
                except Exception:
                    resume_text = raw_bytes.decode("utf-8", errors="ignore")

        elif filename.endswith(".pdf"):
            try:
                pdf_bytes = io.BytesIO(raw_bytes)
                reader = PdfReader(pdf_bytes)
                if len(reader.pages) == 0:
                    return jsonify({"error": "PDF file contains no pages."}), 400

                pages_text = []
                for page in reader.pages:
                    txt = page.extract_text()
                    if txt:
                        pages_text.append(txt)
                resume_text = "\n".join(pages_text)
            except Exception as pdf_err:
                return jsonify({"error": "Failed to parse PDF. File may be corrupted or password-protected."}), 400

        elif filename.endswith(".docx"):
            if docx is None:
                return jsonify({"error": "DOCX parser is not installed on the server."}), 500
            try:
                doc_bytes = io.BytesIO(raw_bytes)
                doc = docx.Document(doc_bytes)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                paragraphs.append(cell.text.strip())
                resume_text = "\n".join(paragraphs)
            except Exception as docx_err:
                return jsonify({"error": "Failed to parse DOCX file. File may be invalid or corrupted."}), 400

        if not resume_text or not resume_text.strip():
            return jsonify({"error": "Uploaded file contains no readable text."}), 400

        clean_text = re.sub(r"[^a-z0-9\s#\+]", " ", resume_text.lower())
        words_set = set(clean_text.split())

        # 1. Skill Detection & Categorization
        detected_skills_set = set()
        categorized_skills = {
            "programming_languages": [],
            "frameworks": [],
            "databases": [],
            "cloud_and_devops": [],
            "tools_and_libraries": [],
            "ai_tools": [],
            "soft_skills": []
        }

        for category_key, category_list in SKILL_CATEGORIES.items():
            for skill in category_list:
                matched = False
                if " " in skill:
                    if skill in clean_text:
                        matched = True
                else:
                    if skill in words_set:
                        matched = True

                if matched:
                    detected_skills_set.add(skill)
                    categorized_skills[category_key].append(skill)

        # 2. Automated Role Prediction
        detected_role = "Software Engineer"
        max_overlap = -1

        for role, expected_list in ROLE_EXPECTED_SKILLS.items():
            overlap_count = 0
            for skill in expected_list:
                if " " in skill:
                    if skill in clean_text:
                        overlap_count += 1
                else:
                    if skill in words_set:
                        overlap_count += 1

            if overlap_count > max_overlap:
                max_overlap = overlap_count
                detected_role = role

        if max_overlap <= 0:
            detected_role = "Software Engineer"

        # 3. Match against Expected Role Benchmark
        expected_skills = ROLE_EXPECTED_SKILLS.get(detected_role, [])
        matched_role_skills = []
        missing_skills = []

        for skill in expected_skills:
            matched = False
            if " " in skill:
                matched = (skill in clean_text)
            else:
                matched = (skill in words_set)

            if matched:
                matched_role_skills.append(skill)
            else:
                missing_skills.append(skill)

        # 4. Section Analysis & Content Extraction
        lines = [line.strip() for line in resume_text.split("\n") if line.strip()]
        
        sections = {
            "summary": any(re.search(r"\b(summary|objective|profile|about)\b", line.lower()) for line in lines),
            "experience": any(re.search(r"\b(experience|employment|work history|career)\b", line.lower()) for line in lines),
            "education": any(re.search(r"\b(education|academic|degree|university|college)\b", line.lower()) for line in lines),
            "projects": any(re.search(r"\b(projects|portfolio|personal projects)\b", line.lower()) for line in lines),
            "certifications": any(re.search(r"\b(certifications|certificates|licenses|credentials)\b", line.lower()) for line in lines)
        }

        # Extract Summary
        summary_text = ""
        for i, line in enumerate(lines[:10]):
            if len(line) > 30 and not re.search(r"@(gmail|yahoo|hotmail|outlook)", line.lower()):
                summary_text = line
                break
        if not summary_text:
            summary_text = lines[0] if lines else "Candidate Resume Profile"

        # Extract Projects highlights
        extracted_projects = []
        in_proj_section = False
        for line in lines:
            if re.search(r"\b(projects|portfolio)\b", line.lower()):
                in_proj_section = True
                continue
            if in_proj_section:
                if any(re.search(rf"\b({s})\b", line.lower()) for s in ["experience", "education", "certifications", "skills"]):
                    in_proj_section = False
                    continue
                if len(line) > 15 and len(extracted_projects) < 5:
                    extracted_projects.append(line)

        if not extracted_projects:
            extracted_projects = [f"Project highlights involving {', '.join(list(detected_skills_set)[:3]) or 'software development'}."]

        # Extract Experience highlights
        extracted_experience = []
        in_exp_section = False
        for line in lines:
            if re.search(r"\b(experience|work history|employment)\b", line.lower()):
                in_exp_section = True
                continue
            if in_exp_section:
                if any(re.search(rf"\b({s})\b", line.lower()) for s in ["education", "projects", "certifications", "skills"]):
                    in_exp_section = False
                    continue
                if len(line) > 20 and len(extracted_experience) < 5:
                    extracted_experience.append(line)

        if not extracted_experience:
            extracted_experience = ["Professional background details detected in resume text."]

        # Extract Education highlights
        extracted_education = []
        for line in lines:
            if any(term in line.lower() for term in ["b.s", "b.a", "bachelor", "master", "m.s", "phd", "university", "degree", "diploma"]):
                extracted_education.append(line)
        if not extracted_education:
            extracted_education = ["Relevant higher education or technical background."]

        # Extract Certifications
        extracted_certs = []
        for line in lines:
            if any(term in line.lower() for term in ["certified", "certification", "coursera", "aws certified", "pmp", "udemy", "license"]):
                extracted_certs.append(line)

        # 5. ATS Score Calculation Logic
        total_benchmark = len(expected_skills)
        matched_count = len(matched_role_skills)
        skill_score = (matched_count / total_benchmark * 45) if total_benchmark > 0 else 20

        present_sections_count = sum(1 for v in sections.values() if v)
        section_score = (present_sections_count / len(sections)) * 30

        content_length = len(resume_text)
        length_score = 15 if 700 <= content_length <= 6000 else 8

        formatting_score = 10 if (sections["experience"] and sections["education"]) else 5

        ats_score = int(round(skill_score + section_score + length_score + formatting_score))
        ats_score = min(100, max(0, ats_score))

        # Strengths & Weaknesses
        strengths = []
        weaknesses = []

        if ats_score >= 70:
            strengths.append(f"Strong keyword alignment for {detected_role} position.")
        else:
            weaknesses.append(f"Low match density for typical {detected_role} requirements.")

        if present_sections_count >= 4:
            strengths.append("Well-structured document with essential ATS sections (Experience, Education, Skills).")
        else:
            weaknesses.append("Missing standard section headers which may hamper ATS document parser indexing.")

        if len(detected_skills_set) >= 8:
            strengths.append(f"Broad technical repertoire identified ({len(detected_skills_set)} total skills detected).")
        else:
            weaknesses.append("Limited technical keywords detected. Add more specific tools, libraries, and frameworks.")

        # Tips
        tips = []
        if missing_skills:
            top_missing = missing_skills[:4]
            tips.append(f"Incorporate core role keywords: {', '.join([s.upper() for s in top_missing])}.")

        if not sections["certifications"]:
            tips.append("Add a 'Certifications' section to highlight formal credentials and online coursework.")

        if not sections["summary"]:
            tips.append("Include a concise 2-3 sentence 'Professional Summary' at the top of your resume.")

        if content_length < 700:
            tips.append("Your resume content is brief. Expand on key responsibilities, deliverables, and measurable metrics.")

        if not tips:
            tips.append("Excellent industry alignment! Boost impact further by quantifying achievements with percentages and numerical metrics.")

        recommendation = f"Candidate shows {ats_score}% alignment for the {detected_role} role. Addressing missing keywords will optimize automated resume screening."

        # Return strict JSON API response
        response_data = {
            "ats_score": ats_score,
            "role": detected_role,
            "summary": summary_text,
            "strengths": strengths,
            "weaknesses": weaknesses,
            "skills": sorted([s.title() for s in detected_skills_set]),
            "skill_categories": {k: sorted([s.title() for s in v]) for k, v in categorized_skills.items()},
            "missing_skills": sorted([s.title() for s in missing_skills]),
            "projects": extracted_projects[:4],
            "experience": extracted_experience[:4],
            "education": extracted_education[:3],
            "certifications": extracted_certs[:3],
            "tips": tips,
            "recommendation": recommendation
        }

        return jsonify(response_data)

    except Exception as e:
        return jsonify({"error": f"An error occurred during analysis: {str(e)}"}), 500

@app.errorhandler(404)
def not_found(e):
    return jsonify({"error": "Resource not found."}), 404

@app.errorhandler(500)
def server_error(e):
    return jsonify({"error": "Internal server error."}), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=PORT)
